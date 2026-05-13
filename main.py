import flet as ft
import os
# Сюда импортируйте ваши функции: universal_parser, create_word_document и т.д.

import tkinter as tk
from tkinter import filedialog
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import traceback
import re

# --- Глобальные настройки стилей ---
FONT_NAME_DOC = 'Times New Roman'
FONT_SIZE_PT_DOC = 14

REPLACEMENTS_MAP = {
    'ќ': 'қ',
    'љ': 'ҷ',
    'ї': 'ӣ',
    'њ': 'ҳ',
}

def replace_specific_chars(text_string):
    if not isinstance(text_string, str):
        return text_string
    for old_char, new_char in REPLACEMENTS_MAP.items():
        text_string = text_string.replace(old_char, new_char)
    return text_string

def read_raw_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.docx':
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return None
    except Exception as e:
        Messagebox.show_error(f"Ошибка при чтении файла '{os.path.basename(filepath)}': {e}\n{traceback.format_exc()}", "Ошибка чтения")
        return None

def universal_parser(filepath):
    raw_text = read_raw_text(filepath)
    if raw_text is None:
        return None

    text = "\n" + raw_text

    MARKER_Q = "___TEMP_Q___"
    MARKER_PLUS = "___TEMP_PLUS___"
    MARKER_MINUS = "___TEMP_MINUS___"

    text = re.sub(r'\n[ \t]*\?', MARKER_Q, text)
    text = re.sub(r'\n[ \t]*\+', MARKER_PLUS, text)
    text = re.sub(r'\n[ \t]*-', MARKER_MINUS, text)

    text = text.replace('#&', MARKER_PLUS)
    text = text.replace('@', MARKER_Q)
    text = text.replace('#', MARKER_MINUS)

    text = text.replace('\n', ' ')

    text = text.replace(MARKER_Q, '\n?')
    text = text.replace(MARKER_PLUS, '\n+')
    text = text.replace(MARKER_MINUS, '\n-')

    tests = []
    current_question_text = None
    current_options_list = []
    current_correct_letter_found = ""
    option_letters = [chr(ord('A') + i) for i in range(26)]

    lines = text.strip().split('\n')
    for line_num, raw_line in enumerate(lines, 1):
        line = raw_line.strip()
        if not line:
            continue
        
        if line.startswith('?'):
            if current_question_text is not None:
                if current_options_list:
                    tests.append({
                        "question": replace_specific_chars(current_question_text),
                        "options": current_options_list,
                        "correct_letter": current_correct_letter_found
                    })
            current_question_text = line[1:].strip()
            current_options_list = []
            current_correct_letter_found = ""
        elif line.startswith('+') or line.startswith('-'):
            if current_question_text is None:
                continue
            is_correct = line.startswith('+')
            option_text = line[1:].strip().rstrip(';')
            if not option_text:
                continue
            
            processed_option_text = replace_specific_chars(option_text)
            current_options_list.append({"text": processed_option_text, "is_correct": is_correct})
            
            if is_correct:
                option_idx = len(current_options_list) - 1
                if option_idx < len(option_letters):
                    current_correct_letter_found = option_letters[option_idx]
                else:
                    current_correct_letter_found = '?'

    if current_question_text is not None and current_options_list:
        if not current_correct_letter_found and any(opt['is_correct'] for opt in current_options_list):
            for i, opt in enumerate(current_options_list):
                if opt['is_correct']:
                    current_correct_letter_found = option_letters[i] if i < len(option_letters) else '?'
                    break
        tests.append({
            "question": replace_specific_chars(current_question_text),
            "options": current_options_list,
            "correct_letter": current_correct_letter_found
        })

    final_tests_validated = []
    for test_item in tests:
        if not test_item["options"]:
            continue
        if not test_item["correct_letter"] and not any(opt['is_correct'] for opt in test_item["options"]):
            test_item["correct_letter"] = ""
        final_tests_validated.append(test_item)

    if not final_tests_validated:
        Messagebox.show_info("Тесты не найдены или все тесты некорректны.\nПроверьте форматы разметки.", "Информация")
        return None
        
    return final_tests_validated

def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'PAGE \* MERGEFORMAT')
    t = OxmlElement('w:t')
    fldSimple.append(t)
    run._r.append(fldSimple)

def create_txt_export(tests_data, output_path):
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            for test in tests_data:
                f.write(f"? {test['question']}\n")
                for opt in test['options']:
                    prefix = "+" if opt['is_correct'] else "-"
                    f.write(f"{prefix} {opt['text']}\n")
                f.write("\n")
        return True
    except Exception as e:
        Messagebox.show_error(f"Не удалось создать TXT файл: {e}\n{traceback.format_exc()}", "Ошибка создания документа")
        return False

def create_word_document(tests_data, output_path, template_path=None):
    try:
        # Если выбран шаблон, открываем его и добавляем новый раздел (новую страницу)
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            # Иначе создаем пустой документ
            doc = Document()
            section = doc.sections[0]

        # Настраиваем стиль шрифта по умолчанию для документа
        style = doc.styles['Normal']
        font = style.font
        font.name = FONT_NAME_DOC
        font.size = Pt(FONT_SIZE_PT_DOC)

        # Настраиваем отступы ТОЛЬКО для текущего раздела (чтобы не сломать титульник)
        section.top_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.right_margin = Cm(1.2)
        section.header_distance = Cm(0.0)
        section.footer_distance = Cm(0.5)
        section.different_first_page_header_footer = False

        # Настройка нумерации страниц
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), '2') # Страница с тестами всегда будет номером 2

        # Отвязываем колонтитул от предыдущего раздела (титульника), чтобы номер страницы не появился на титульнике
        footer = section.footer
        if template_path:
            footer.is_linked_to_previous = False
            
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = ""
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_field(footer_para)

        # Создание таблицы
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(0.74)
        table.columns[1].width = Cm(18.14)

        hdr_cells = table.add_row().cells
        hdr_cells[0].text = ''
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_hdr0 = hdr_cells[0].paragraphs[0].runs[0] if hdr_cells[0].paragraphs[0].runs else hdr_cells[0].paragraphs[0].add_run()
        run_hdr0.font.name = FONT_NAME_DOC
        run_hdr0.font.size = Pt(FONT_SIZE_PT_DOC)

        hdr_cells[1].text = 'Вопрос и варианты ответов'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_hdr1 = hdr_cells[1].paragraphs[0].runs[0]
        run_hdr1.font.name = FONT_NAME_DOC
        run_hdr1.font.size = Pt(FONT_SIZE_PT_DOC)

        for test_idx, test_item in enumerate(tests_data):
            row_cells = table.add_row().cells
            cell1 = row_cells[0]
            p1 = cell1.paragraphs[0] if cell1.paragraphs else cell1.add_paragraph()
            p1.text = ""
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            correct_letter_to_display = test_item.get("correct_letter", "")
            if not correct_letter_to_display and test_item.get("options"):
                correct_letter_to_display = "$$"
            run1 = p1.add_run(correct_letter_to_display)
            run1.font.name = FONT_NAME_DOC
            run1.font.size = Pt(FONT_SIZE_PT_DOC)

            cell2 = row_cells[1]
            p_question = cell2.paragraphs[0] if cell2.paragraphs else cell2.add_paragraph()
            p_question.text = ""
            
            run_q_prefix = p_question.add_run(f"{test_idx + 1}. @")
            run_q_prefix.font.name = FONT_NAME_DOC
            run_q_prefix.font.size = Pt(FONT_SIZE_PT_DOC)
            run_q_text = p_question.add_run(" " + test_item["question"])
            run_q_text.font.name = FONT_NAME_DOC
            run_q_text.font.size = Pt(FONT_SIZE_PT_DOC)
            run_q_text.bold = True
            
            for opt_idx, option_data in enumerate(test_item["options"]):
                p_option = cell2.add_paragraph()
                paragraph_format = p_option.paragraph_format
                tab_stops = paragraph_format.tab_stops
                if not any(ts.position == Cm(0.5) for ts in tab_stops):
                    tab_stops.add_tab_stop(Cm(0.5), WD_TAB_ALIGNMENT.LEFT)
                p_option.add_run('\t')
                run_opt_letter = p_option.add_run(f"{chr(ord('A') + opt_idx)}) ")
                run_opt_letter.font.name = FONT_NAME_DOC
                run_opt_letter.font.size = Pt(FONT_SIZE_PT_DOC)
                run_opt_text = p_option.add_run(option_data['text'])
                run_opt_text.font.name = FONT_NAME_DOC
                run_opt_text.font.size = Pt(FONT_SIZE_PT_DOC)
        
        doc.save(output_path)
        return True
    except Exception as e:
        Messagebox.show_error(f"Не удалось создать Word документ: {e}\n{traceback.format_exc()}", "Ошибка создания документа")
        return False

class TestConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.geometry("450x250") # Немного увеличил окно
        self.template_path = None # Переменная для хранения пути к шаблону

        main_frame = ttkb.Frame(self.root, padding=(20, 15))
        main_frame.pack(expand=True, fill=BOTH)

        title_label = ttkb.Label(main_frame, text="Конвертер Тестов", font=("-family {Segoe UI} -size 16 -weight bold"))
        title_label.pack(pady=(0, 15))

        # --- Блок выбора шаблона ---
        template_frame = ttkb.Frame(main_frame)
        template_frame.pack(fill=X, pady=5)
        
        self.lbl_template_info = ttkb.Label(template_frame, text="Шаблон титульного: Не выбран", foreground="gray")
        self.lbl_template_info.pack(side=TOP, anchor=W, pady=(0, 5))
        
        btn_template = ttkb.Button(template_frame, text="1. Выбрать титульный лист (Опционально)", command=self.choose_template, bootstyle=(INFO, OUTLINE))
        btn_template.pack(fill=X)

        # --- Основная кнопка ---
        btn_process = ttkb.Button(main_frame, text="2. Выбрать тесты и Конвертировать", command=self.process_file, bootstyle=(PRIMARY, SOLID))
        btn_process.pack(fill=X, pady=15, ipady=8)
        
        self.root.eval('tk::PlaceWindow . center')

    def choose_template(self):
        filepath = filedialog.askopenfilename(
            title="Выберите Word-файл с титульным листом",
            filetypes=(("Word Документы", "*.docx"), ("Все файлы", "*.*"))
        )
        if filepath:
            self.template_path = filepath
            filename = os.path.basename(filepath)
            self.lbl_template_info.config(text=f"Шаблон: {filename}", foreground="green")

    def process_file(self):
        input_file_path = filedialog.askopenfilename(
            title="Выберите файл с тестами (DOCX или TXT)",
            filetypes=(("Все поддерживаемые форматы", "*.docx *.txt"), ("Word Документы", "*.docx"), ("Текстовые файлы", "*.txt"), ("Все файлы", "*.*"))
        )
        if not input_file_path:
            Messagebox.show_info("Файл с тестами не выбран. Операция отменена.", "Операция отменена")
            return

        tests_data = universal_parser(input_file_path)
        if not tests_data:
            return

        base_name = os.path.basename(input_file_path)
        file_name_without_ext = os.path.splitext(base_name)[0]
        default_output_docx_name = file_name_without_ext + "_готовый.docx"
        initial_dir = os.path.dirname(input_file_path)

        docx_output_path = filedialog.asksaveasfilename(
            title="Сохранить ГОТОВЫЙ DOCX файл как...",
            initialfile=default_output_docx_name,
            defaultextension=".docx",
            initialdir=initial_dir,
            filetypes=(("Word Документы", "*.docx"), ("Все файлы", "*.*"))
        )
        if not docx_output_path:
            Messagebox.show_info("Место для сохранения не выбрано. Операция отменена.", "Операция отменена")
            return

        base_output_dir = os.path.dirname(docx_output_path)
        base_output_name = os.path.splitext(os.path.basename(docx_output_path))[0]
        txt_output_path = os.path.join(base_output_dir, base_output_name + ".txt")

        # Передаем путь к шаблону (если он был выбран)
        doc_success = create_word_document(tests_data, docx_output_path, self.template_path)
        txt_success = create_txt_export(tests_data, txt_output_path)

        if doc_success and txt_success:
            Messagebox.show_success(
                f"Файлы успешно обработаны и сохранены:\n\n1. {os.path.basename(docx_output_path)}\n2. {os.path.basename(txt_output_path)}", 
                "Успех!"
            )
            self.show_placeholder_warnings(tests_data)
            self.root.quit()

    def show_placeholder_warnings(self, tests_data):
        questions_with_placeholder = []
        for idx, test_item in enumerate(tests_data):
            if not test_item.get("correct_letter", "") and test_item.get("options"):
                questions_with_placeholder.append(f"   - Вопрос {idx + 1}: \"{test_item['question'][:40]}...\"")
        
        if questions_with_placeholder:
            warning_title = "Предупреждение о содержании"
            intro_text = "Внимание! Для следующих вопросов не был указан правильный ответ (отмеченный знаком '+').\n"
            
            full_warning_text = intro_text
            full_warning_text += f"В документе для них проставлен символ '$$':\n\n"
            full_warning_text += "\n".join(questions_with_placeholder)
            Messagebox.show_warning(full_warning_text, warning_title)


if __name__ == "__main__":
    selected_theme = "cosmo" 
    app_root = ttkb.Window(themename=selected_theme, title="Конвертер Тестов")
    app = TestConverterApp(app_root)
    app_root.mainloop()

def main(page: ft.Page):
    page.title = "Конвертер Тестов"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.window_width = 400
    page.window_height = 700
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    # Переменные для хранения путей
    template_path = None
    test_path = None

    # UI элементы
    status_text = ft.Text("Готов к работе", color=ft.colors.GREY)
    
    # --- Логика FilePicker ---
    def on_template_picked(e: ft.FilePickerResultEvent):
        nonlocal template_path
        if e.files:
            template_path = e.files[0].path
            btn_template.text = f"Шаблон: {e.files[0].name}"
            btn_template.icon = ft.icons.CHECK_CIRCLE
            page.update()

    def on_test_picked(e: ft.FilePickerResultEvent):
        nonlocal test_path
        if e.files:
            test_path = e.files[0].path
            btn_test.text = f"Файл: {e.files[0].name}"
            btn_test.icon = ft.icons.CHECK_CIRCLE
            page.update()

    def on_save_picked(e: ft.FilePickerResultEvent):
        if e.path and test_path:
            status_text.value = "Обработка..."
            page.update()
            
            try:
                # ЗДЕСЬ ВАША ЛОГИКА
                # tests_data = universal_parser(test_path)
                # create_word_document(tests_data, e.path, template_path)
                # create_txt_export(tests_data, e.path.replace(".docx", ".txt"))
                
                status_text.value = f"Успешно сохранено!"
                status_text.color = ft.colors.GREEN
            except Exception as ex:
                status_text.value = f"Ошибка: {ex}"
                status_text.color = ft.colors.RED
            page.update()

    # Инициализация диалогов выбора файлов
    pick_template_dialog = ft.FilePicker(on_result=on_template_picked)
    pick_test_dialog = ft.FilePicker(on_result=on_test_picked)
    save_file_dialog = ft.FilePicker(on_result=on_save_picked)
    
    page.overlay.extend([pick_template_dialog, pick_test_dialog, save_file_dialog])

    # --- Кнопки ---
    btn_template = ft.ElevatedButton(
        "1. Выбрать титульный (Опционально)",
        icon=ft.icons.UPLOAD_FILE,
        on_click=lambda _: pick_template_dialog.pick_files(allowed_extensions=["docx"])
    )

    btn_test = ft.ElevatedButton(
        "2. Выбрать тесты",
        icon=ft.icons.DESCRIPTION,
        on_click=lambda _: pick_test_dialog.pick_files(allowed_extensions=["docx", "txt"])
    )

    btn_process = ft.FilledButton(
        "3. Конвертировать и Сохранить",
        icon=ft.icons.SAVE,
        on_click=lambda _: save_file_dialog.save_file(
            file_name="Готовые_тесты.docx",
            allowed_extensions=["docx"]
        )
    )

    # Собираем интерфейс
    page.add(
        ft.Text("Конвертер Тестов", size=24, weight=ft.FontWeight.BOLD),
        ft.Divider(),
        btn_template,
        btn_test,
        ft.Container(height=20),
        btn_process,
        ft.Container(height=20),
        status_text
    )

ft.app(target=main)